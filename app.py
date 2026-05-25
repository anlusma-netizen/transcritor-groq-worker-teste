import os
import re
import json
import shutil
import tempfile
import time
import subprocess
from pathlib import Path
from typing import Optional, List, Dict, Any
from urllib.parse import urlparse, parse_qs

import requests
import gdown
from fastapi import FastAPI, UploadFile, File, Request, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from groq import Groq

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


VERSION = "20.0.0"

app = FastAPI(title="Worker Telegram → Groq → DOCX", version=VERSION)

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")
GROQ_TRANSCRIPTION_MODEL = os.getenv("GROQ_TRANSCRIPTION_MODEL", "whisper-large-v3-turbo")
GROQ_TRANSLATION_MODEL = os.getenv("GROQ_TRANSLATION_MODEL", "llama-3.3-70b-versatile")
RAW_TRANSLATION_MODEL = os.getenv("RAW_TRANSLATION_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct")
TRANSLATE_NON_PT = os.getenv("TRANSLATE_NON_PT", "true").lower() in {"1", "true", "yes", "sim"}
PARAGRAPH_SENTENCES = int(os.getenv("PARAGRAPH_SENTENCES", "5"))
PARAGRAPH_MAX_CHARS = int(os.getenv("PARAGRAPH_MAX_CHARS", "1400"))
MAX_AUDIO_MB = int(os.getenv("MAX_AUDIO_MB", "24"))
TARGET_AUDIO_BITRATE = os.getenv("TARGET_AUDIO_BITRATE", "24k")
TARGET_AUDIO_FORMAT = os.getenv("TARGET_AUDIO_FORMAT", "mp3")
TRANSLATION_CHUNK_CHARS = int(os.getenv("TRANSLATION_CHUNK_CHARS", "3000"))
TRANSLATION_DELAY_SECONDS = int(os.getenv("TRANSLATION_DELAY_SECONDS", "12"))
TRANSLATION_RETRY_WAIT_SECONDS = int(os.getenv("TRANSLATION_RETRY_WAIT_SECONDS", "70"))
INCLUDE_ORIGINAL_FOR_NON_PT = os.getenv("INCLUDE_ORIGINAL_FOR_NON_PT", "true").lower() in {"1", "true", "yes", "sim"}

client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None


@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    return JSONResponse(
        status_code=500,
        content={
            "ok": False,
            "version": VERSION,
            "error": str(exc),
            "type": exc.__class__.__name__,
        },
    )


@app.get("/")
def root():
    return {
        "ok": True,
        "service": "transcritor-groq-worker-teste",
        "version": VERSION,
        "output_format": "docx",
        "copy_structure": "raw_paragraphs",
        "fast_mode": true,
        "raw_copy_mode": true,
        "translation_style": "ptbr_copy_analysis",
        "translate_non_pt": TRANSLATE_NON_PT,
        "raw_translation_model": RAW_TRANSLATION_MODEL,
        "routes": ["/health", "/process-source", "/process-telegram-media"],
    }


@app.get("/health")
def health():
    return {
        "ok": True,
        "version": VERSION,
        "groq_key_configured": bool(GROQ_API_KEY),
        "telegram_token_configured": bool(TELEGRAM_BOT_TOKEN),
        "transcription_model": GROQ_TRANSCRIPTION_MODEL,
        "translation_model": GROQ_TRANSLATION_MODEL,
        "translation_chunk_chars": TRANSLATION_CHUNK_CHARS,
        "translation_delay_seconds": TRANSLATION_DELAY_SECONDS,
        "output_format": "docx",
        "copy_structure": "raw_paragraphs",
        "fast_mode": true,
        "raw_copy_mode": true,
        "translation_style": "ptbr_copy_analysis",
        "translate_non_pt": TRANSLATE_NON_PT,
        "raw_translation_model": RAW_TRANSLATION_MODEL,
        "include_original_for_non_pt": INCLUDE_ORIGINAL_FOR_NON_PT,
    }


def run_cmd(cmd: List[str]):
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"Erro no comando: {' '.join(cmd)}\nSTDERR:\n{result.stderr}")
    return result


def safe_filename(name: str, fallback: str = "arquivo") -> str:
    name = name or fallback
    name = re.sub(r"[^\w\-. ]+", "_", name, flags=re.UNICODE).strip()
    return name[:120] or fallback


def download_url(url: str, output_path: Path):
    with requests.get(url, stream=True, timeout=180) as r:
        r.raise_for_status()
        with open(output_path, "wb") as f:
            for chunk in r.iter_content(chunk_size=1024 * 1024):
                if chunk:
                    f.write(chunk)


def extract_google_drive_file_id(url: str) -> Optional[str]:
    if "drive.google.com" not in url:
        return None

    m = re.search(r"/file/d/([^/]+)", url)
    if m:
        return m.group(1)

    qs = parse_qs(urlparse(url).query)
    if "id" in qs:
        return qs["id"][0]

    return None


def download_google_drive_or_url(url: str, output_path: Path):
    file_id = extract_google_drive_file_id(url)

    if file_id:
        gdown_url = f"https://drive.google.com/uc?id={file_id}"
        result = gdown.download(gdown_url, str(output_path), quiet=False, fuzzy=True)
        if not result or not output_path.exists() or output_path.stat().st_size < 1024:
            raise RuntimeError(
                "Não consegui baixar o arquivo do Google Drive. "
                "Confirme que está em 'Qualquer pessoa com o link pode visualizar'."
            )
        return

    download_url(url, output_path)
    if not output_path.exists() or output_path.stat().st_size < 1024:
        raise RuntimeError("Download falhou ou retornou arquivo vazio.")


def download_telegram_file(file_id: str, output_path: Path):
    if not TELEGRAM_BOT_TOKEN:
        raise RuntimeError("TELEGRAM_BOT_TOKEN não configurado no Railway de teste.")

    info_url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/getFile"
    info = requests.get(info_url, params={"file_id": file_id}, timeout=60).json()

    if not info.get("ok"):
        raise RuntimeError(f"Telegram getFile falhou: {info}")

    file_path = info["result"]["file_path"]
    file_url = f"https://api.telegram.org/file/bot{TELEGRAM_BOT_TOKEN}/{file_path}"
    download_url(file_url, output_path)


def convert_to_audio(input_path: Path, output_path: Path):
    cmd = [
        "ffmpeg", "-y",
        "-i", str(input_path),
        "-vn",
        "-ac", "1",
        "-ar", "16000",
        "-b:a", TARGET_AUDIO_BITRATE,
        str(output_path),
    ]
    run_cmd(cmd)


def file_mb(path: Path) -> float:
    return path.stat().st_size / (1024 * 1024)


def split_audio(input_audio: Path, chunks_dir: Path, chunk_seconds: int = 600) -> List[Path]:
    chunks_dir.mkdir(parents=True, exist_ok=True)
    pattern = chunks_dir / f"chunk_%03d.{TARGET_AUDIO_FORMAT}"
    cmd = [
        "ffmpeg", "-y",
        "-i", str(input_audio),
        "-f", "segment",
        "-segment_time", str(chunk_seconds),
        "-c", "copy",
        str(pattern),
    ]
    run_cmd(cmd)
    return sorted(chunks_dir.glob(f"chunk_*.{TARGET_AUDIO_FORMAT}"))


def transcribe_one(audio_path: Path) -> Dict[str, Any]:
    if not client:
        raise RuntimeError("GROQ_API_KEY não configurada no Railway.")

    with open(audio_path, "rb") as f:
        result = client.audio.transcriptions.create(
            file=(audio_path.name, f.read()),
            model=GROQ_TRANSCRIPTION_MODEL,
            response_format="verbose_json",
            temperature=0,
        )

    if hasattr(result, "model_dump"):
        return result.model_dump()

    if isinstance(result, dict):
        return result

    return json.loads(result.json())


def transcribe_audio(audio_path: Path, workdir: Path) -> Dict[str, Any]:
    if file_mb(audio_path) <= MAX_AUDIO_MB:
        return transcribe_one(audio_path)

    chunks = split_audio(audio_path, workdir / "chunks", chunk_seconds=600)
    all_text = []
    language = None

    for index, chunk in enumerate(chunks, start=1):
        print(f"Transcrevendo chunk {index}/{len(chunks)}: {chunk.name}")
        res = transcribe_one(chunk)

        if not language:
            language = res.get("language")

        text = res.get("text", "")
        if text:
            all_text.append(text.strip())

    return {
        "text": "\n\n".join(all_text).strip(),
        "language": language,
    }


def is_portuguese_language(language: Optional[str]) -> bool:
    if not language:
        return False
    lang = str(language).lower().strip()
    return lang in {"pt", "por", "portuguese", "português", "portugues", "pt-br", "pt_br"} or "portugu" in lang


def split_text_for_llm(text: str, max_chars: int) -> List[str]:
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    parts = []
    current = []
    current_len = 0

    for paragraph in paragraphs:
        chunks = [paragraph[i:i + max_chars] for i in range(0, len(paragraph), max_chars)] if len(paragraph) > max_chars else [paragraph]

        for chunk in chunks:
            if current and current_len + len(chunk) + 1 > max_chars:
                parts.append("\n".join(current).strip())
                current = []
                current_len = 0

            current.append(chunk)
            current_len += len(chunk) + 1

    if current:
        parts.append("\n".join(current).strip())

    return [p for p in parts if p]


def run_llm_with_retry(prompt: str, index: int, total: int) -> str:
    if not client:
        raise RuntimeError("GROQ_API_KEY não configurada no Railway.")

    last_error = None

    for attempt in range(1, 6):
        try:
            completion = client.chat.completions.create(
                model=GROQ_TRANSLATION_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": "Você organiza transcrições de anúncios, VSLs e cartas de vendas com fidelidade, sem reescrever criativamente.",
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0,
            )
            return completion.choices[0].message.content.strip()

        except Exception as exc:
            last_error = exc
            error_text = str(exc)

            if (
                "rate_limit_exceeded" in error_text
                or "Request too large" in error_text
                or "tokens per minute" in error_text
                or "TPM" in error_text
                or "Rate limit" in error_text
            ):
                print(f"Rate limit no bloco {index}/{total}. Tentativa {attempt}/5. Esperando {TRANSLATION_RETRY_WAIT_SECONDS}s.")
                time.sleep(TRANSLATION_RETRY_WAIT_SECONDS)
                continue

            raise

    raise RuntimeError(f"Falha ao processar bloco {index}/{total}: {last_error}")




def clean_transcript_text(text: str) -> str:
    text = (text or "").strip()
    # Remove timestamps comuns se aparecerem: [00:01], 00:01, 00:01:20
    text = re.sub(r"\[?\b\d{1,2}:\d{2}(?::\d{2})?\b\]?", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def split_sentences(text: str) -> List[str]:
    text = clean_transcript_text(text)
    if not text:
        return []
    # Divide em frases, mas sem inventar estrutura de copy.
    sentences = re.split(r"(?<=[.!?])\s+", text)
    return [s.strip() for s in sentences if s.strip()]


def paragraphize_raw_copy(text: str, sentences_per_paragraph: Optional[int] = None, max_chars: Optional[int] = None) -> List[str]:
    """
    Só diagrama a cópia em parágrafos legíveis.
    Não cria Hook/Body/CTA.
    Não resume.
    Não interpreta.
    """
    sentences_per_paragraph = sentences_per_paragraph or PARAGRAPH_SENTENCES
    max_chars = max_chars or PARAGRAPH_MAX_CHARS

    sentences = split_sentences(text)
    if not sentences:
        cleaned = clean_transcript_text(text)
        return [cleaned] if cleaned else ["[transcrição vazia]"]

    paragraphs = []
    current = []

    for sentence in sentences:
        candidate = " ".join(current + [sentence]).strip()

        if current and (len(current) >= sentences_per_paragraph or len(candidate) > max_chars):
            paragraphs.append(" ".join(current).strip())
            current = [sentence]
        else:
            current.append(sentence)

    if current:
        paragraphs.append(" ".join(current).strip())

    return [p for p in paragraphs if p.strip()]


def run_translation_with_retry(prompt: str, index: int, total: int, model: str) -> str:
    if not client:
        raise RuntimeError("GROQ_API_KEY não configurada no Railway.")

    last_error = None

    for attempt in range(1, 4):
        try:
            completion = client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Você é um tradutor/adaptador profissional para copywriters brasileiros. "
                            "Sua tarefa é transformar uma copy em inglês em uma versão em português brasileiro natural, "
                            "boa para análise interna de copy, preservando intenção, agressividade, dados, promessas e ordem das ideias. "
                            "Não censure, não suavize, não resuma, não moralize e não omita termos sensíveis."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0,
            )
            return completion.choices[0].message.content.strip()

        except Exception as exc:
            last_error = exc
            error_text = str(exc)

            if (
                "rate_limit_exceeded" in error_text
                or "Request too large" in error_text
                or "tokens per minute" in error_text
                or "TPM" in error_text
                or "Rate limit" in error_text
                or "429" in error_text
            ):
                wait_time = min(25, max(5, TRANSLATION_RETRY_WAIT_SECONDS // 3))
                print(f"Rate limit na tradução {index}/{total}. Tentativa {attempt}/3. Esperando {wait_time}s.")
                time.sleep(wait_time)
                continue

            # fallback para modelo principal, caso o modelo escolhido não exista ou falhe
            if model != GROQ_TRANSLATION_MODEL:
                print(f"Modelo de tradução falhou ({model}). Tentando fallback: {GROQ_TRANSLATION_MODEL}")
                return run_translation_with_retry(prompt, index, total, GROQ_TRANSLATION_MODEL)

            raise

    raise RuntimeError(f"Falha ao traduzir bloco {index}/{total} com {model}: {last_error}")


def translate_raw_copy_to_ptbr(text: str) -> str:
    text = clean_transcript_text(text)
    if not text:
        return ""

    parts = split_text_for_llm(text, TRANSLATION_CHUNK_CHARS)
    translated = []

    for i, part in enumerate(parts, start=1):
        prompt = f"""
Traduza/adapte a transcrição abaixo para português brasileiro natural, como material de análise interna para um copywriter brasileiro.

Objetivo:
- A tradução deve soar como uma copy entendível para brasileiros, não como tradução literal.
- Preserve a intenção da copy original, a agressividade, o exagero, o tom provocativo, a promessa, a ordem das ideias e os dados.
- A versão em português deve ajudar o copywriter a entender o mecanismo persuasivo da copy.

Regras obrigatórias:
- Não resuma.
- Não analise.
- Não adicione títulos.
- Não separe em Hook, Body, CTA, promessa ou oferta.
- Não censure, não suavize, não moralize e não omita palavras sensíveis, palavrões, termos sexuais ou promessas agressivas.
- Não traduza literalmente expressões que fiquem estranhas em português brasileiro.
- Adapte expressões idiomáticas para equivalentes naturais no Brasil.
- Preserve nomes, marcas, números, provas, promessas e sequência das ideias.
- Quando houver medidas em polegadas, mantenha a referência original e acrescente a equivalência aproximada em centímetros quando ficar natural. Exemplo: 9 inches → 9 polegadas, cerca de 22 cm.
- Mantenha termos de marketing em inglês quando forem comuns para copywriter brasileiro: hook, CTA, VSL, lead, offer, pitch, upsell, funnel, checkout.
- Evite frases robóticas como “o tamanho vai disparar”, “atores de ponta”, “ela vai agradecer por assistir”. Prefira português brasileiro natural e direto.
- Entregue apenas a tradução corrida em parágrafos.

Transcrição original:
{part}
""".strip()

        translated.append(run_translation_with_retry(prompt, i, len(parts), RAW_TRANSLATION_MODEL))

    return "\n\n".join(translated).strip()


def prepare_structured_text(text: str, source_is_portuguese: bool) -> str:
    """
    v19: cópia crua diagramada.
    Sem Hook/Body/CTA.
    Sem mapa de VSL.
    Só parágrafos.
    """
    base_text = clean_transcript_text(text)
    if not base_text:
        return "MODE: RAW\n\nCÓPIA DIAGRAMADA:\n[transcrição vazia]"

    if source_is_portuguese or not TRANSLATE_NON_PT:
        working_text = base_text
        label = "CÓPIA DIAGRAMADA"
    else:
        working_text = translate_raw_copy_to_ptbr(base_text)
        label = "CÓPIA TRADUZIDA PARA ANÁLISE EM PT-BR"

    paragraphs = paragraphize_raw_copy(working_text)

    return (
        "MODE: RAW\n\n"
        + f"{label}:\n"
        + "\n\n".join(paragraphs)
    )



def parse_raw_sections(structured_text: str) -> Dict[str, List[str]]:
    sections: Dict[str, List[str]] = {}
    current = None

    for raw_line in structured_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            if current:
                sections[current].append("")
            continue

        if stripped.startswith("MODE:"):
            continue

        if stripped.endswith(":"):
            current = stripped[:-1].strip()
            sections.setdefault(current, [])
            continue

        if current is None:
            current = "CÓPIA DIAGRAMADA"
            sections.setdefault(current, [])

        sections[current].append(stripped)

    return sections


def setup_docx_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Arial"

    styles["Title"].font.size = Pt(20)
    styles["Title"].font.bold = True

    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)


def add_paragraphs(doc: Document, lines: List[str]):
    buffer = []

    def flush():
        if not buffer:
            return
        paragraph_text = " ".join([x.strip() for x in buffer if x.strip()]).strip()
        buffer.clear()
        if paragraph_text:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.space_after = Pt(9)
            p.paragraph_format.line_spacing = 1.18
            p.add_run(paragraph_text)

    for line in lines:
        if not line.strip():
            flush()
        else:
            buffer.append(line)

    flush()


def create_docx(original_name: str, transcription: Dict[str, Any], structured_text: str, output_path: Path, include_original: bool):
    doc = Document()
    setup_docx_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    source_is_pt = is_portuguese_language(transcription.get("language"))

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if source_is_pt:
        title.add_run("Cópia Crua Diagramada")
    else:
        title.add_run("Cópia Traduzida para Análise PT-BR")

    meta = doc.add_paragraph(style="Normal")
    meta.add_run("Arquivo: ").bold = True
    meta.add_run(original_name)

    meta2 = doc.add_paragraph(style="Normal")
    meta2.add_run("Idioma detectado: ").bold = True
    meta2.add_run(str(transcription.get("language") or "não identificado"))

    meta3 = doc.add_paragraph(style="Normal")
    meta3.add_run("Formato: ").bold = True
    meta3.add_run("cópia crua em parágrafos; tradução adaptada para análise, sem Hook/Body/CTA")

    doc.add_paragraph("")

    sections = parse_raw_sections(structured_text)

    for heading, lines in sections.items():
        doc.add_heading(heading, level=1)
        add_paragraphs(doc, lines)

    if include_original:
        original_text = transcription.get("text", "") or ""
        original_paragraphs = paragraphize_raw_copy(original_text)
        if original_text.strip():
            doc.add_page_break()
            doc.add_heading("TRANSCRIÇÃO ORIGINAL", level=1)
            add_paragraphs(doc, sum(([p, ""] for p in original_paragraphs), []))

    doc.save(output_path)


def process_file(input_path: Path, original_name: str, workdir: Path) -> Path:
    audio_path = workdir / f"audio_convertido.{TARGET_AUDIO_FORMAT}"
    convert_to_audio(input_path, audio_path)

    transcription = transcribe_audio(audio_path, workdir)
    source_is_portuguese = is_portuguese_language(transcription.get("language"))

    structured_text = prepare_structured_text(
        transcription.get("text", ""),
        source_is_portuguese=source_is_portuguese,
    )

    include_original = INCLUDE_ORIGINAL_FOR_NON_PT and not source_is_portuguese

    output_name = safe_filename(Path(original_name).stem or "transcricao") + "_copia_analise_ptbr.docx"
    output_path = workdir / output_name
    create_docx(original_name, transcription, structured_text, output_path, include_original=include_original)
    return output_path


def persistent_file_response(output_path: Path) -> FileResponse:
    final_path = Path(tempfile.gettempdir()) / output_path.name
    shutil.copy(output_path, final_path)
    return FileResponse(
        final_path,
        filename=output_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.post("/process-telegram-media")
async def process_telegram_media(request: Request, file: Optional[UploadFile] = File(default=None)):
    with tempfile.TemporaryDirectory() as tmp:
        workdir = Path(tmp)

        if file is None:
            body = await request.body()
            if not body:
                raise HTTPException(status_code=400, detail="Nenhum arquivo recebido.")
            original_name = request.headers.get("x-file-name", "arquivo_telegram")
            input_path = workdir / safe_filename(original_name, "input.bin")
            input_path.write_bytes(body)
        else:
            original_name = request.headers.get("x-file-name", file.filename or "arquivo_telegram")
            input_path = workdir / safe_filename(original_name, "input.bin")
            with open(input_path, "wb") as f:
                f.write(await file.read())

        output_path = process_file(input_path, original_name, workdir)
        return persistent_file_response(output_path)


@app.post("/process-source")
async def process_source(payload: Dict[str, Any]):
    source_type = payload.get("sourceType")
    original_name = payload.get("fileName") or "arquivo"

    with tempfile.TemporaryDirectory() as tmp:
        workdir = Path(tmp)
        input_path = workdir / safe_filename(original_name, "input.bin")

        if source_type == "url":
            url = payload.get("url")
            if not url:
                raise HTTPException(status_code=400, detail="URL não enviada.")
            download_google_drive_or_url(url, input_path)

        elif source_type == "telegram_file_id":
            file_id = payload.get("fileId")
            if not file_id:
                raise HTTPException(status_code=400, detail="fileId do Telegram não enviado.")
            download_telegram_file(file_id, input_path)

        else:
            raise HTTPException(status_code=400, detail="Envie arquivo ou link público.")

        output_path = process_file(input_path, original_name, workdir)
        return persistent_file_response(output_path)
